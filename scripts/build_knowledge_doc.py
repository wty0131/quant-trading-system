#!/usr/bin/env python3
"""
构建量化交易系统知识手册 Word 文档
输出到桌面: 量化交易系统知识手册.docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 样式设置 ──
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = '微软雅黑'
    heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_colored_table(doc, headers, rows, col_widths=None):
    """创建带表头的彩色表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    return table


# ═══════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('量化交易系统\n知识手册')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('v1.0 — 六阶段完整系统\n2026年6月')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ═══════════════════════════════════════════
# 目录
# ═══════════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '一、核心指标速查表 — 变量清单与标准值',
    '二、系统思维导图',
    '三、数据层 — 多市场统一接口',
    '四、回测引擎 — 事件驱动架构',
    '五、策略库 — 六个经典策略',
    '六、风控与组合管理',
    '七、执行层 — 从回测到交易',
    '八、通往实盘 — 用真实持股做量化',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# ═══════════════════════════════════════════
# 一、核心指标速查表
# ═══════════════════════════════════════════
doc.add_heading('一、核心指标速查表', level=1)

doc.add_paragraph(
    '以下表格列出量化交易系统的全部核心变量、计算公式、标准阈值及使用说明。'
)

# 表1：收益与风险指标
doc.add_heading('1.1 收益与风险指标', level=2)

headers = ['指标', '公式', '优秀', '合格', '差', '如何比较']
rows = [
    ['年化收益率', '日均收益 × 252', '> 15%', '5~15%', '< 0%',
     '与无风险利率(2.5%)比：>5%才有超额价值。与基准(沪深300)比：超过基准才算Alpha'],
    ['年化波动率', '日std × √252', '< 15%', '15~25%', '> 30%',
     '越低越稳。但低波动≠好策略——需配合收益看Sharpe。加密>50%正常'],
    ['夏普比率\n(Sharpe)', '(R年化-Rf)/σ年化', '> 1.0', '0.5~1.0', '< 0',
     '每单位风险的超额收益。>1不错，>2优秀。<0说明跑不赢定存'],
    ['索提诺比率\n(Sortino)', '(R年化-Rf)/σ下行', '> 1.5', '1.0~1.5', '< 0.5',
     '只看下行波动。比Sharpe更真实——上涨波动不该被惩罚'],
    ['最大回撤\n(MDD)', 'min((P-Peak)/Peak)', '< -10%', '-10~-25%', '< -30%',
     '绝对值越小越好。实盘MDD通常比回测大30%。与收益比=Calmar'],
    ['卡尔玛比率\n(Calmar)', 'R年化/|MDD|', '> 1.0', '0.5~1.0', '< 0.3',
     '每单位最大回撤的年化回报。>1=一年赚回最惨亏损'],
    ['Beta', 'Cov(Rs,Rm)/Var(Rm)', '0.3~0.7', '0.7~1.3', '>1.5',
     '对基准敏感度。<1防御型，>1进攻型。<0.3才有分散效果'],
    ['Alpha', 'Rs均值-β×Rm均值', '> 0.05', '0~0.05', '< 0',
     '剔除市场涨跌后的纯超额收益。Jensen Alpha年化>5%才算有真本事'],
]
add_colored_table(doc, headers, rows)

doc.add_paragraph()
doc.add_paragraph(
    '比较逻辑链：先看Sharpe(有没有超额收益) → 再看MDD(最坏亏多少) → '
    '再看Calmar(收益是否值得承受这个回撤) → 最后看Beta(是不是只是β赌方向)。'
).runs[0].bold = True

# 表2：交易统计指标
doc.add_heading('1.2 交易统计指标', level=2)

headers2 = ['指标', '公式', '优秀', '合格', '差', '如何比较']
rows2 = [
    ['胜率', '盈利交易/总交易', '45~55%', '35~45%', '< 30%',
     '不是越高越好。趋势跟踪胜率30~40%正常。高胜率+低盈亏比=假象'],
    ['盈亏比\n(ProfitFactor)', '总盈利/总亏损', '> 2.0', '1.3~2.0', '< 1.0',
     '核心指标。>1.3才值得做。胜率×盈亏比必须>1才有正期望'],
    ['平均盈亏', 'avg(盈利)/avg(亏损)', '> 2.0', '1.5~2.0', '< 1.0',
     '平均每笔赚的÷每笔亏的。与胜率互补——胜率低但盈亏比高=海龟'],
    ['持仓占比', '持仓天数/总天数', '50~70%', '30~50%', '< 20%',
     '策略"在市场中"的时间。太低=择时太激进，太高≈买入持有'],
    ['信息比率(IR)', 'IC均值/IC波动', '> 0.5', '0.3~0.5', '< 0.2',
     '因子/策略超额收益的稳定性。>0.5=每次都跑赢，0.2=偶尔跑赢'],
    ['最大连胜/连败', '—', '连败<8', '连败<15', '连败>20',
     '心理承受力测试。连败超过你能承受的=你会手动干预策略'],
]
add_colored_table(doc, headers2, rows2)

doc.add_paragraph()
doc.add_paragraph(
    '判断一个策略好坏的三步法：① 交易次数>30(统计显著) → '
    '② 盈亏比>1.3 + 胜率>30%(正期望) → '
    '③ Sharpe>0.5 + Calmar>0.5(风险调整后仍为正) → 值得考虑'
).runs[0].bold = True

# 表3：组合与分配指标
doc.add_heading('1.3 组合与风控指标', level=2)

headers3 = ['指标', '标准值', '说明']
rows3 = [
    ['策略相关性阈值', '< 0.5 可配对，< 0.3 好',
     '相关>0.5的组合效果不明显。还要看回撤重叠度——平时相关低≠危机时不一起跌'],
    ['Kelly 仓位', 'f* = (p×b-(1-p))/b，实际用 half-Kelly',
     '全Kelly太激进。p=0.4, b=2→f=10%，half→5%。单笔不得超过25%'],
    ['日亏损上限', '≤ 5%',
     '触发后只平不开，次日重置。超出=策略已失控'],
    ['最大回撤上限', '≤ 20%',
     '触发后全部清仓。这是最后防线——比策略止损更硬'],
    ['单品种仓位上限', '≤ 30%',
     '一只股票不超过总资金30%。即使看对，黑天鹅也会干死集中持仓'],
    ['总仓位上限', '≤ 80%',
     '永远留20%现金。不是保守——是给加仓和保证金留余地'],
    ['波动率倒数权重', 'w_i = (1/vol_i)/Σ(1/vol_j)',
     '最稳健的资金分配法。不碰收益预测，只依赖波动率持续性'],
    ['样本外验证', '训练集:验证集 = 70:30 或滚动窗口',
     '回测训练=在样本内调参。必须留30%数据从没见过=真正检验'],
]
add_colored_table(doc, headers3, rows3)

doc.add_page_break()

# ═══════════════════════════════════════════
# 二、思维导图
# ═══════════════════════════════════════════
doc.add_heading('二、系统思维导图', level=1)

mind_map = """
┌──────────────────────────────────────────────────────────────────┐
│                     量化交易系统 v1.0                              │
├──────────────────────────────────────────────────────────────────┤
│                                                                  │
│  ┌──────────┐    ┌──────────┐    ┌──────────┐    ┌─────────────┐ │
│  │  数据层   │───▶│ 回测引擎  │───▶│  策略库   │───▶│ 风控+组合   │ │
│  │ (data/)  │    │(backtest/)│    │(strateg/) │    │  (risk/)    │ │
│  └────┬─────┘    └────┬─────┘    └────┬─────┘    └──────┬──────┘ │
│       │               │               │                  │        │
│  · 三市场         · 事件驱动       · 双均线           · 仓位管理   │
│    统一接口          MarketEvent      布林带             Kelly     │
│    A股/加密/美股     SignalEvent      海龟系统           风险平价   │
│  · SQLite          OrderEvent        RSRS              · 止损系统  │
│    WAL+UPSERT      FillEvent         多因子              ATR/移动  │
│  · 统一Schema     · Engine.run()    配对交易            · 策略相关  │
│    10列标准        · 买入持有验证    · 回测验证          · 资金分配  │
│                                    │                  │        │
│                                    ▼                  ▼        │
│                              ┌──────────────────────────┐       │
│                              │     执行层 (execution/)    │       │
│                              │  Broker + OMS + RiskGuard │       │
│                              │  PaperEngine → 纸交易     │       │
│                              └────────────┬─────────────┘       │
│                                           │                     │
│                                           ▼                     │
│                              ┌──────────────────────────┐       │
│                              │  仪表盘 (dashboard/)       │       │
│                              │  Streamlit 4页纯Python    │       │
│                              │  总览|策略|回测|风控       │       │
│                              └──────────────────────────┘       │
│                                                                  │
│  ════════════════════ 学习路径 ════════════════════               │
│                                                                  │
│  N00: 手工数据探索 ──→ N01: 数据管道 ──→ N02: 回测引擎            │
│        ↓                    ↓                ↓                   │
│   Sharpe/MDD/MA       ABC抽象层       事件驱动 vs 向量化          │
│                                                                  │
│  N03: 策略库 ──→ N04: 风控组合 ──→ N05: 执行层                   │
│        ↓              ↓              ↓                           │
│   6个策略回测      仓位+止损+分配   Broker+OMS+纸交易             │
│                                                                  │
│  实盘路径: 纸交易2周 → 偏差分析 → QMT开通 → 小资金试 → 逐步加码   │
│                                                                  │
└──────────────────────────────────────────────────────────────────┘
"""
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(mind_map)
run.font.name = 'Consolas'
run.font.size = Pt(7.5)

doc.add_page_break()

# ═══════════════════════════════════════════
# 三、数据层
# ═══════════════════════════════════════════
doc.add_heading('三、数据层 — 多市场统一接口', level=1)

doc.add_heading('3.1 ABC 抽象基类模式', level=2)
doc.add_paragraph(
    '核心思想：定义契约(DataSource)，子类各自实现 _fetch()，框架自动完成批量拉取+标准化。'
)
doc.add_paragraph(
    'DataSource(ABC) → AShareSource(baostock直连) / CryptoSource(ccxt自动切换Gate.io/Kraken) / '
    'USStockSource(yfinance+SOCKS5代理)。三个市场用同一个 get_history() 方法。'
)

doc.add_heading('3.2 统一 Schema', level=2)
doc.add_paragraph(
    '10列标准：[symbol, date, open, high, low, close, volume, amount, market, interval]。'
    '所有市场输出完全一致的列——这是系统的数据契约。'
)

doc.add_heading('3.3 SQLite 存储', level=2)
doc.add_paragraph(
    'WAL模式：允许多进程同时读写，不锁表。(见 N01 第6章)\n'
    'UPSERT语义：同(symbol, date)重复写入不翻倍——增量更新安全。(见 N01 第6章)\n'
    '三层分离：API → DataFrame(内存可检查) → SQLite(磁盘缓存)。不是"API直接灌SQL"。'
)

doc.add_heading('3.4 三市场数据源', level=2)
doc.add_paragraph('A股: baostock (独立源，国内直连，无需Token) → 见 N01 第4章')
doc.add_paragraph('加密: Gate.io直连 + Kraken/Binance via SOCKS5代理 → 见 N01 第5章')
doc.add_paragraph('美股: yfinance + v2rayN SOCKS5 (127.0.0.1:10808) → 见 N01 第6章')

doc.add_page_break()

# ═══════════════════════════════════════════
# 四、回测引擎
# ═══════════════════════════════════════════
doc.add_heading('四、回测引擎 — 事件驱动架构', level=1)

doc.add_heading('4.1 为什么事件驱动？', level=2)
doc.add_paragraph(
    '向量化(Notebook 00的双均线)有两个死穴：① 前瞻偏差——rolling(20)在t=19能看到t=20的数据；'
    '② 无法处理路径依赖——止损、限价单、加仓都依赖"当前状态"。(见 N02 第1章)'
)
doc.add_paragraph(
    '事件驱动逐条回放历史数据：每个t时刻只知道t及以前的信息。与实盘逻辑完全一致。'
)

doc.add_heading('4.2 事件流转', level=2)
doc.add_paragraph(
    'MarketEvent(新K线) → Strategy.on_bar() → SignalEvent(开/平仓) → '
    'Portfolio.generate_order() → OrderEvent → ExecutionHandler → '
    'FillEvent → Portfolio.update() → 更新持仓+净值。(见 N02 第2-6章)'
)

doc.add_heading('4.3 引擎验证', level=2)
doc.add_paragraph(
    '买入持有策略：回测结果与手动计算误差 < 0.017%（零摩擦模式下）。'
    '这是回测引擎正确性的基准校验。(见 N02 第6章)'
)

doc.add_heading('4.4 绩效分析模块', level=2)
doc.add_paragraph(
    'backtest/analytics.py 输出BacktestReport：总收益/年化收益/Sharpe/Sortino/MDD/Calmar/胜率/盈亏比/持仓占比。'
    '每个指标的计算公式和标准值 → 参见本手册"一、核心指标速查表"。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 五、策略库
# ═══════════════════════════════════════════
doc.add_heading('五、策略库 — 六个经典策略', level=1)

doc.add_paragraph(
    '策略 = on_bar() 里的决策规则。你的工作是设计规则，引擎负责验证。\n'
    '每个策略的详细原理 → 见 N03 对应章节。'
)

strategies = [
    ('5.1 双均线 (DualMA)', 'MA5上穿MA20=金叉买入，下穿=死叉卖出', '趋势跟踪 ★',
     '信号滞后性认知。震荡市中反复止损是所有趋势跟踪的共同特征。见 N03 第1章'),
    ('5.2 布林带 (Bollinger)', '价格<下轨(MA-2σ)=买入，价格>中轨=平仓', '均值回归 ★★',
     '赌统计规律"95%时间价格在带内"。震荡市好用，趋势市被打爆。带宽=波动率指标。见 N03 第2章'),
    ('5.3 海龟交易 (Turtle)', '突破20日高=入场，ATR×2=止损，每0.5ATR加仓，破10日低=出场', '完整系统 ★★',
     '胜率极低(~30%)但盈亏比极高。95%时间亏小钱，5%时间赚大钱。见 N03 第3章'),
    ('5.4 RSRS', 'H~L做OLS回归，beta斜率标准化z-score与阈值比较', '阻力支撑 ★★',
     '用数学替代主观画线。beta大=买方强。见 N03 第4章'),
    ('5.5 多因子选股', '动量+反转+低波+成交量→综合得分→买前N名→定期调仓', '截面Alpha ★★★',
     '横截面策略(哪只更好)vs时间序列策略(何时买卖)。核心检验：因子IC。见 N03 第5章'),
    ('5.6 配对交易', '两只高相关股票的价差spread做z-score→>2做空价差→回归平仓', '统计套利 ★★★★',
     '需要先验证协整关系(价差必须均值回归)。对冲比率beta的控制。见 N03 第6章'),
]

for title, logic, diff, detail in strategies:
    doc.add_heading(title, level=2)
    doc.add_paragraph(f'逻辑：{logic}')
    doc.add_paragraph(f'难度：{diff}')
    doc.add_paragraph(f'要点：{detail}')

doc.add_page_break()

# ═══════════════════════════════════════════
# 六、风控与组合管理
# ═══════════════════════════════════════════
doc.add_heading('六、风控与组合管理', level=1)

doc.add_heading('6.1 仓位管理', level=2)
doc.add_paragraph(
    '四种模型：(1)固定比例——每次用X%资金；(2)Kelly公式——f*=(p×b-(1-p))/b，实战用half-Kelly；'
    '(3)风险平价——高波动买少低波动买多；(4)波动率目标——动态调整维持目标波动。见 N04 第2章'
)

doc.add_heading('6.2 止损系统', level=2)
doc.add_paragraph(
    '在on_bar()中：止损 > 止盈 > 正常信号。止损优先级最高——先保命再赚钱。\n'
    '四种止损：固定(-N%)/ATR(动态跟随)/移动(N日最低)/时间(超时不涨)。见 N04 第3章'
)

doc.add_heading('6.3 策略相关性', level=2)
doc.add_paragraph(
    '日收益相关性 < 0.5 = 可配对。但相关性只看线性关系，看不到"尾部相关"——'
    '两个策略在平时相关低，但黑天鹅日可能同时暴跌。还要检查回撤重叠度。< 30%重叠才算真正的好配对。见 N04 第4章'
)

doc.add_heading('6.4 资金分配', level=2)
doc.add_paragraph(
    '等权=基准(无知假设)。波动率倒数=最稳健(只依赖波动率持续性，不碰收益预测)。'
    'Max Sharpe=理论最优但容易过拟合(过去收益≠未来收益)。见 N04 第5章'
)

doc.add_heading('6.5 组合验证 (CSI300 2024)', level=2)
doc.add_paragraph(
    '4策略组合结果：组合Sharpe 0.570 > 单策略均值0.445。组合MDD -5.85% < 所有单策略最低 -6.28%。'
    '资金分配(波动率倒数)：Bollinger 43% / DualMA 21% / Turtle 20% / RSRS 16%。'
    '这个结果证明：分散化是量化中为数不多的"免费午餐"。见 N04 第6章'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 七、执行层
# ═══════════════════════════════════════════
doc.add_heading('七、执行层 — 从回测到交易', level=1)

doc.add_heading('7.1 回测 vs 实盘的鸿沟', level=2)
doc.add_paragraph(
    '回测假设：看到价格=能成交，立即成交，无滑点，无限流动性。\n'
    '实盘真相：订单排在队列里，可能部分成交／被拒，价格继续变，你的大单会推动价格。\n'
    '执行层管理这些"回测不会告诉你"的摩擦。见 N05 第1章'
)

doc.add_heading('7.2 Broker 统一接口', level=2)
doc.add_paragraph(
    'PaperBroker: 模拟成交，立刻可用。CCXTBroker: 加密实盘(Gate.io)。QMTBroker: A股实盘(待券商开通)。\n'
    '同一个抽象模型，纸交易↓实盘切换。见 N05 第2章'
)

doc.add_heading('7.3 订单状态机', level=2)
doc.add_paragraph(
    'submit → PENDING → (PARTIAL) → FILLED/CANCELLED。订单有独立的生命周期——不像回测那样立即成交。\n'
    '超时重试、部分成交剩余撤单、异常降级——OMS管理这一切。见 N05 第3章'
)

doc.add_heading('7.4 纸交易引擎', level=2)
doc.add_paragraph(
    'PaperTradingEngine = BacktestEngine的现实版。相同结构(strategy.on_bar→signal→order→fill)，'
    '不同数据源(实时baostock/ccxt代替SQLite历史数据)。见 N05 第4章'
)

doc.add_heading('7.5 独立风控守护', level=2)
doc.add_paragraph(
    '策略代码不能裁判自己。RiskGuard独立运行：日亏损>5%→只平不开，最大回撤>20%→全部清仓，'
    '单品种>30%→限该品种，总仓位>80%→限所有。见 N05 第5章'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 八、通往实盘
# ═══════════════════════════════════════════
doc.add_heading('八、通往实盘 — 用真实持股做量化交易', level=1)

doc.add_heading('8.1 你现在用的是模拟数据吗？', level=2)
doc.add_paragraph(
    '你的系统已经用的是真实数据：baostock 拉取的是沪深300真实行情（1323条日线，2020-2025），'
    '回测结果（海龟 10.72%、Sharpe 0.668）都是基于真实价格计算。\n\n'
    '"模拟"指的是资金和成交——你的策略发出买入信号，引擎假设你能按bar.close成交，'
    '没有人真的在交易所替你下单。这就是纸交易(Paper Trading)的定位。'
)

doc.add_heading('8.2 从纸交易到真实持股的路径', level=2)

steps = [
    ('第1步：纸交易验证 (当前立即可做)',
     '启动纸交易引擎 → 每天收盘后拉取baostock最新日线 → 策略on_bar()决策 → PaperBroker记录虚拟成交 → '
     '积累2-4周真实时间线的交易记录。\n'
     '目的：验证(1)策略规则在真实时间线上能执行 (2)你不会在连亏后手动干预 (3)信号频率是否合理。\n'
     '工具：execution/paper_engine.py → 手动每天跑一次，或加APScheduler定时。'),

    ('第2步：偏差分析',
     '纸交易2-4周后，对比纸交易净值 vs 回测同一时段的净值。\n'
     '偏差来源：(1)滑点模型不准(实盘滑点更大) (2)你的执行延迟 (3)信号不连续。\n'
     '根据偏差调整execution的slippage和commission_rate参数，重新回测。'),

    ('第3步：开通QMT/A股实盘接口',
     '联系你的A股券商，申请开通QMT(迅投)量化交易权限。\n'
     '开通后：安装xtquant包 → 修改qmt_broker.py中的_connect()逻辑 → 将PaperBroker替换为QMTBroker。\n'
     'QMTBroker的框架代码已就绪(execution/qmt_broker.py)，只需填入券商给你的连接参数。'),

    ('第4步：小资金试水',
     '实盘第一周只用回测资金的10%(如10万)。\n'
     '观察：(1)订单是否正常成交 (2)滑点与回测模型偏差 (3)系统稳定性(不崩)。\n'
     '对比实盘成交价 vs 纸交易模拟价，校准ExecutionHandler。'),

    ('第5步：逐步加码',
     '连续2周无异常 → 资金加到30%。\n'
     '连续1月无异常 → 资金加到50%。\n'
     '从不加到100%——永远留20%现金备用。\n'
     '同时启用RiskGuard的日亏损上限和最大回撤清仓。'),
]

for title, desc in steps:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_heading('8.3 加密实盘（可同步进行）', level=2)
doc.add_paragraph(
    '如果你有Gate.io账户：创建API Key → 写入.env → CCXTBroker(testnet=False) → '
    '先跑testnet(测试网，免费假币) → 确认OK → 切到现货实盘。\n'
    '加密市场24/7交易，没有涨跌停/T+1限制，是验证策略执行层的最佳场所。\n'
    'ccxt_broker.py已就绪，SOCKS5代理自动走v2rayN 10808端口。'
)

doc.add_heading('8.4 你的真实持股怎么接入？', level=2)
doc.add_paragraph(
    '如果你持有某几只A股（如贵州茅台、招商银行）：\n'
    '1) 用 baostock 拉取这些股票的历史日线 → store.save("ashare", "daily")\n'
    '2) 在 Dashboard 回测页选择这些股票 → 跑你的策略 → 看历史回测结果\n'
    '3) 如果回测Sharpe>0.5且MDD<15% → 考虑用该策略管理这些持仓\n'
    '4) 用纸交易引擎模拟该策略的每日操作 → 对比你现在的持仓决策\n'
    '5) 开通QMT后，策略可以自动执行：当策略发出卖出信号时自动下单\n\n'
    '注意：量化策略不应该完全取代你的判断——它应该是你决策的"第二意见"。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 附录：Notebook索引
# ═══════════════════════════════════════════
doc.add_heading('附录：Notebook 教学索引', level=1)

notebooks = [
    ('N00', 'notebooks/00_data_exploration.ipynb', '数据探索',
     '用真实A股数据手算Sharpe/MDD/双均线。纯手工，理解底层。'),
    ('N01', 'notebooks/01_data_pipeline.ipynb', '数据管道',
     'ABC抽象基类 → 三市场统一接口 → SQLite存储。从手工到自动化。'),
    ('N02', 'notebooks/02_backtest_engine.ipynb', '回测引擎',
     '事件驱动 → 事件类型 → Strategy基类 → Portfolio → 引擎组装。'),
    ('N03', 'notebooks/03_strategies.ipynb', '策略库',
     '6个策略逐个实现：布林带/海龟/RSRS/多因子/配对。回测对比汇总。'),
    ('N04', 'notebooks/04_risk_portfolio.ipynb', '风控组合',
     '仓位管理/止损/策略相关性/资金分配/组合回测。'),
    ('N05', 'notebooks/05_live_execution.ipynb', '执行层',
     'Broker/OMS/PaperEngine/RiskGuard。从回测到交易。'),
]

for nid, path, title, desc in notebooks:
    doc.add_heading(f'{nid}: {title}', level=2)
    doc.add_paragraph(f'文件: {path}')
    doc.add_paragraph(f'内容: {desc}')

# ── 保存 ──
output_path = os.path.expanduser('~/Desktop/量化交易系统知识手册.docx')
doc.save(output_path)
print(f'文档已保存到: {output_path}')
